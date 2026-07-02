# -*- coding: utf-8 -*-
"""Tạo PHIẾU THÔNG TIN để client điền → tuỳ chỉnh các file quy tắc AI của CTRC."""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Mm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = r"C:\Users\Admin\Downloads\CTRC_Skill_Rules\CTRC - Phieu thong tin chinh sua (client dien).docx"
ACCENT = RGBColor(0xE9,0x4E,0x1B)
GREY   = RGBColor(0x88,0x88,0x88)
FILLBG = "FFF8E1"   # vàng nhạt = ô điền
HEADBG = "FCE3D3"   # cam nhạt = header bảng

CRIT = [
 ("🎯 Tập trung","Lắng nghe, duy trì sự chú ý và theo kịp hoạt động học tập",
  ["Tập trung tốt trong giờ học.","Lắng nghe hướng dẫn cẩn thận.","Theo sát hoạt động của lớp.","Chú ý quan sát và phản hồi nhanh.","Duy trì sự tập trung tốt."],
  ["Đôi lúc còn mất tập trung trong hoạt động.","Cần thêm nhắc nhở để duy trì sự chú ý.","Thỉnh thoảng bị phân tâm bởi môi trường xung quanh.","Chưa duy trì được sự tập trung xuyên suốt buổi học.","Cần hỗ trợ để quay lại nhiệm vụ khi bị gián đoạn."]),
 ("🙋 Tham gia","Hứng thú, chủ động tham gia trò chơi, hoạt động và thảo luận",
  ["Tham gia hoạt động tích cực.","Hào hứng với các thử thách.","Chủ động phát biểu ý kiến.","Mạnh dạn tham gia trò chơi.","Tương tác tốt cùng cô và bạn."],
  ["Còn khá dè dặt khi tham gia hoạt động.","Chưa chủ động chia sẻ ý kiến của mình.","Cần khuyến khích thêm để tham gia thảo luận.","Tham gia hoạt động khi có sự động viên từ cô.","Chưa thực sự tự tin khi thể hiện bản thân."]),
 ("🧠 Tư duy","Quan sát, phân loại, nhận biết quy luật, suy luận và giải quyết vấn đề phù hợp độ tuổi",
  ["Quan sát tốt.","Nhận biết quy luật nhanh.","Phân loại chính xác.","Suy luận phù hợp độ tuổi.","Biết thử nhiều cách giải.","Xử lý nhiệm vụ linh hoạt."],
  ["Cần thêm thời gian để tìm ra cách giải quyết.","Còn gặp khó khăn khi nhận biết quy luật mới.","Chưa mạnh dạn thử các hướng giải khác nhau.","Cần thêm gợi ý để hoàn thành nhiệm vụ.","Đang từng bước làm quen với dạng bài mới."]),
 ("🌱 Tự lập","Tự thực hiện nhiệm vụ, tự tìm cách giải quyết trước khi nhờ hỗ trợ",
  ["Hoàn thành nhiệm vụ độc lập.","Chủ động thực hiện yêu cầu.","Tự tin khi làm bài.","Có tinh thần tự giác tốt.","Kiên trì hoàn thành nhiệm vụ."],
  ["Thường tìm sự hỗ trợ trước khi tự thử.","Cần nhắc nhở để bắt đầu nhiệm vụ.","Chưa thực sự tự tin khi làm việc độc lập.","Cần hỗ trợ thêm trong quá trình thực hiện nhiệm vụ.","Đôi lúc bỏ cuộc khi gặp thử thách."]),
 ("🤝 Hợp tác","Tương tác, làm việc cùng bạn, chờ đến lượt và chia sẻ ý kiến",
  ["Hợp tác tốt cùng bạn bè.","Tương tác tích cực trong nhóm.","Biết chờ đến lượt.","Chia sẻ và hỗ trợ bạn.","Hòa đồng với các bạn.","Tham gia nhóm rất tích cực."],
  ["Còn ít tương tác với các bạn trong nhóm.","Cần khuyến khích thêm khi tham gia hoạt động nhóm.","Chưa chủ động trao đổi cùng bạn bè.","Đôi lúc gặp khó khăn khi phối hợp với nhóm.","Cần thêm thời gian để hòa nhập với hoạt động chung."]),
]

doc = Document()
sec = doc.sections[0]
sec.page_width=Mm(210); sec.page_height=Mm(297)
sec.left_margin=sec.right_margin=Mm(18); sec.top_margin=sec.bottom_margin=Mm(16)
doc.styles['Normal'].font.name='Calibri'; doc.styles['Normal'].font.size=Pt(11)

def shade(cell, fill):
    tcPr=cell._tc.get_or_add_tcPr(); shd=OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:fill'),fill); tcPr.append(shd)
def setw(cell,w): cell.width=w
def fillbox(label, hint='', widthlabel=Inches(2.4)):
    t=doc.add_table(rows=1, cols=2); t.autofit=False
    L,R=t.rows[0].cells
    L.width=widthlabel; R.width=Inches(4.6)
    lp=L.paragraphs[0]; r=lp.add_run(label); r.bold=True
    if hint:
        hp=L.add_paragraph(); hr=hp.add_run(hint); hr.italic=True; hr.font.size=Pt(8.5); hr.font.color.rgb=GREY
    shade(R, FILLBG); R.paragraphs[0].add_run(' ')
    doc.add_paragraph()
def choices(label, opts, hint=''):
    p=doc.add_paragraph(); r=p.add_run(label); r.bold=True
    if hint:
        hp=doc.add_paragraph(); hr=hp.add_run(hint); hr.italic=True; hr.font.size=Pt(8.5); hr.font.color.rgb=GREY
    cp=doc.add_paragraph()
    cp.add_run('   '.join('☐ '+o for o in opts))
def h1(t):
    h=doc.add_heading('',level=1); r=h.add_run(t); r.font.color.rgb=ACCENT
def h2(t): doc.add_heading(t,level=2)
def note(t):
    p=doc.add_paragraph(); r=p.add_run(t); r.italic=True; r.font.size=Pt(9.5); r.font.color.rgb=GREY

# ===== TITLE =====
ti=doc.add_heading('',level=0); ti.add_run('PHIẾU THÔNG TIN TUỲ CHỈNH BÁO CÁO AI — CTRC')
st=doc.add_paragraph(); sr=st.add_run('WonderKids · Quý khách điền phiếu này để hệ thống AI viết báo cáo ĐÚNG Ý trung tâm.'); sr.italic=True
doc.add_paragraph('Điền trực tiếp vào các ô màu vàng / tích vào ☐. Phần nào muốn GIỮ NGUYÊN thì bỏ trống. Hoàn tất gửi lại cho bộ phận kỹ thuật để cập nhật vào hệ thống (có hiệu lực ngay).')
nb=doc.add_paragraph(); nbr=nb.add_run('Lưu ý: quý khách chỉnh CÁCH VIẾT / GIỌNG / NỘI DUNG thoải mái. KHÔNG cần đụng tới phần kỹ thuật (tên trường JSON) — bộ phận kỹ thuật lo phần đó.')
nbr.bold=True; nbr.font.color.rgb=ACCENT
doc.add_paragraph('Người điền: ……………………………………   ·   Chức vụ: ………………………   ·   Ngày: ……/……/………')

# ===== A. THƯƠNG HIỆU =====
h1('A. Thông tin trung tâm & thương hiệu')
fillbox('Tên trung tâm (hiện đầu báo cáo)','VD: WonderKids Edu')
fillbox('Slogan / tagline (nếu có)','VD: Khai mở tư duy')
fillbox('Màu thương hiệu','mã màu hoặc mô tả, VD: cam #E94E1B')
fillbox('Câu kết đặt cuối mỗi báo cáo','Mặc định: "Mỗi bé có tốc độ riêng — với sự đồng hành kiên nhẫn, bé sẽ ngày càng tự tin và phát triển toàn diện."')

# ===== B. GIỌNG VĂN =====
h1('B. Giọng văn gửi phụ huynh')
note('Tương ứng file: 01 - Giọng thương hiệu')
choices('Tông giọng mong muốn:', ['Ấm áp, gần gũi','Chuyên nghiệp, súc tích','Trang trọng','Vui tươi'], 'tick 1–2 ô; ghi rõ thêm ở ô dưới nếu cần')
fillbox('Mô tả thêm về giọng','VD: ấm áp nhưng vẫn cho thấy chuyên môn')
fillbox('Cách xưng hô với trẻ','VD: "con" / "bé" / "học sinh"')
fillbox('Cách gọi giáo viên','VD: "cô" / "thầy" / "giáo viên"')
fillbox('Từ ngữ ĐẶC TRƯNG muốn dùng','các từ/cụm thể hiện thương hiệu, mỗi từ cách nhau dấu phẩy')
fillbox('Từ CẤM dùng (thêm ngoài mặc định)','Mặc định đã cấm: kém, yếu, tệ, dốt, lười, chậm chạp, ngu, hư')
choices('Dùng emoji trong tin nhắn?', ['Có, nhẹ nhàng','Không dùng'], '')
fillbox('Ví dụ 1 câu nhận xét "đúng giọng" mà trung tâm thích','để AI học theo văn phong này')

# ===== C. 5 TIÊU CHÍ + CÂU MẪU =====
h1('C. 5 tiêu chí đánh giá & ngân hàng câu mẫu')
note('Tương ứng file: 02 - Báo cáo per-buổi. Mỗi buổi giáo viên tick 1 điểm mạnh + 1 điều cần cải thiện; AI dựa câu mẫu để viết. Dưới đây là bản HIỆN TẠI — quý khách sửa chữ, xoá câu không thích, hoặc thêm câu mới vào dòng trống.')
for label,desc,pos,neg in CRIT:
    h2(label)
    fillbox('Đổi tên / mô tả tiêu chí (nếu muốn)','Hiện tại: '+desc, widthlabel=Inches(2.4))
    t=doc.add_table(rows=1, cols=2); t.style='Light Grid Accent 2'; t.autofit=False
    a,b=t.rows[0].cells; a.width=Inches(3.5); b.width=Inches(3.5)
    for c,txt in ((a,'CÂU TÍCH CỰC (điểm mạnh)'),(b,'CÂU CẦN CẢI THIỆN')):
        c.paragraphs[0].clear(); rr=c.paragraphs[0].add_run(txt); rr.bold=True; shade(c,HEADBG)
    rows=max(len(pos),len(neg))+2   # +2 dòng trống để thêm câu mới
    for i in range(rows):
        cells=t.add_row().cells; cells[0].width=Inches(3.5); cells[1].width=Inches(3.5)
        cells[0].paragraphs[0].add_run(pos[i] if i<len(pos) else '')
        cells[1].paragraphs[0].add_run(neg[i] if i<len(neg) else '')
    doc.add_paragraph()

# ===== D. BÁO CÁO PER-BUỔI =====
h1('D. Cấu trúc báo cáo từng buổi')
note('Tương ứng file: 02 - Báo cáo per-buổi. Hiện gồm 3 phần: 📖 Nội dung học tuần · ✅ Điểm bé làm được · 📌 Điểm bé làm chưa được.')
choices('Giữ mục "📖 Nội dung học tuần này" ở đầu báo cáo?', ['Giữ','Bỏ'], '')
fillbox('Đổi nhãn "✅ Điểm bé làm được" (nếu muốn)','để trống = giữ nguyên')
fillbox('Đổi nhãn "📌 Điểm bé làm chưa được" (nếu muốn)','để trống = giữ nguyên')
fillbox('Độ dài mong muốn mỗi phần','VD: 1–2 câu / ngắn gọn / chi tiết hơn')
fillbox('Yêu cầu khác cho báo cáo buổi','')

# ===== E. BÁO CÁO TUẦN =====
h1('E. Cấu trúc báo cáo tuần')
note('Tương ứng file: 03 - Báo cáo tuần. Hiện gồm: 📖 Tuần này con học · 🌟 Tổng hợp tuần · 📌 Điểm cần cải thiện · 💬 Lời nhắn cô.')
choices('Có cần thêm mục "Tiến bộ so với tuần trước"?', ['Có','Không'], '')
fillbox('Thêm / bớt / đổi nhãn các mục báo cáo tuần','')
fillbox('Độ dài & yêu cầu khác cho báo cáo tuần','')

# ===== F. CÁCH GỬI =====
h1('F. Cách gửi phụ huynh')
note('Tương ứng file: gửi-phụ-huynh. Nguyên tắc: KHÔNG gửi link lạ.')
choices('Hình thức gửi:', ['Văn bản (copy Zalo)','Ảnh báo cáo','Cả hai'], '')
fillbox('Yêu cầu khác khi gửi','')

# ===== G. GÓP Ý =====
h1('G. Góp ý / yêu cầu khác')
t=doc.add_table(rows=1,cols=1); c=t.rows[0].cells[0]; shade(c,FILLBG)
for _ in range(4): c.add_paragraph(' ')

doc.add_paragraph()
foot=doc.add_paragraph('Gửi phiếu đã điền về: ……………………………………  ·  CTRC — WonderKids · phiên bản quy tắc 2026-06-09')
foot.runs[0].font.size=Pt(8.5); foot.runs[0].font.color.rgb=GREY

os.makedirs(os.path.dirname(OUT), exist_ok=True)
doc.save(OUT)
print('wrote', OUT)
