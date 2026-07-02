# -*- coding: utf-8 -*-
"""Chuyển các file rule .md (CTRC) -> .docx sạch đẹp để gửi khách hàng chỉnh sửa."""
import os, re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

RULES = r"C:\Users\Admin\Downloads\wks\ctrc\rules"
OUT   = r"C:\Users\Admin\Downloads\CTRC_Skill_Rules"
os.makedirs(OUT, exist_ok=True)

FILES = [
    ("giong-thuong-hieu.md",  "01 - Giong thuong hieu"),
    ("bao-cao-per-buoi.md",   "02 - Bao cao per-buoi"),
    ("bao-cao-tuan.md",       "03 - Bao cao tuan"),
    ("doc-hieu-giao-an.md",   "04 - Doc hieu giao an"),
]

ACCENT = RGBColor(0xE9,0x4E,0x1B)   # cam thương hiệu

def add_inline(p, text):
    """Parse **bold** và `code` thành runs."""
    # tách theo **...** và `...`
    parts = re.split(r'(\*\*.+?\*\*|`.+?`)', text)
    for seg in parts:
        if not seg:
            continue
        if seg.startswith('**') and seg.endswith('**'):
            r = p.add_run(seg[2:-2]); r.bold = True
        elif seg.startswith('`') and seg.endswith('`'):
            r = p.add_run(seg[1:-1]); r.font.name = 'Consolas'; r.font.size = Pt(10)
        else:
            p.add_run(seg)

def render_md(doc, md, as_section=False, section_title=None):
    lines = md.split('\n')
    i = 0
    n = len(lines)
    while i < n:
        line = lines[i].rstrip('\n')
        s = line.strip()
        # bỏ comment metadata
        if s.startswith('<!--'):
            i += 1; continue
        # bảng markdown
        if s.startswith('|') and i+1 < n and re.match(r'^\|[\s:|-]+\|?\s*$', lines[i+1].strip()):
            rows = []
            while i < n and lines[i].strip().startswith('|'):
                rows.append(lines[i].strip()); i += 1
            header = [c.strip() for c in rows[0].strip('|').split('|')]
            data = []
            for r in rows[2:]:
                data.append([c.strip() for c in r.strip('|').split('|')])
            t = doc.add_table(rows=1, cols=len(header))
            t.style = 'Light Grid Accent 2'
            for j,h in enumerate(header):
                cell = t.rows[0].cells[j]; cell.paragraphs[0].clear()
                add_inline(cell.paragraphs[0], h)
                for run in cell.paragraphs[0].runs: run.bold = True
            for drow in data:
                cells = t.add_row().cells
                for j,val in enumerate(drow):
                    if j < len(cells):
                        cells[j].paragraphs[0].clear(); add_inline(cells[j].paragraphs[0], val)
            doc.add_paragraph()
            continue
        # code block ```
        if s.startswith('```'):
            i += 1; code=[]
            while i < n and not lines[i].strip().startswith('```'):
                code.append(lines[i]); i += 1
            i += 1
            p = doc.add_paragraph()
            r = p.add_run('\n'.join(code)); r.font.name='Consolas'; r.font.size=Pt(9.5)
            p.paragraph_format.left_indent = Inches(0.2)
            continue
        # headings
        if s.startswith('### '):
            doc.add_heading(s[4:], level=3); i+=1; continue
        if s.startswith('## '):
            doc.add_heading(s[3:], level=2); i+=1; continue
        if s.startswith('# '):
            if as_section:
                doc.add_heading(s[2:], level=1)
            else:
                h = doc.add_heading('', level=0); add_inline(h, s[2:])
            i+=1; continue
        # blockquote
        if s.startswith('> '):
            p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(0.25)
            r0 = p.add_run('▸ '); r0.font.color.rgb = ACCENT; r0.bold=True
            add_inline(p, s[2:]);
            for run in p.runs: run.italic = True
            i+=1; continue
        # bullet list
        m = re.match(r'^[-*]\s+(.*)$', s)
        if m:
            p = doc.add_paragraph(style='List Bullet'); add_inline(p, m.group(1)); i+=1; continue
        # numbered list
        m = re.match(r'^(\d+)\.\s+(.*)$', s)
        if m:
            p = doc.add_paragraph(style='List Number'); add_inline(p, m.group(2)); i+=1; continue
        # blank
        if s == '':
            i+=1; continue
        # normal paragraph
        p = doc.add_paragraph(); add_inline(p, s); i+=1

def base_doc():
    doc = Document()
    st = doc.styles['Normal']; st.font.name='Calibri'; st.font.size=Pt(11)
    return doc

# ---- file lẻ ----
for fn, title in FILES:
    md = open(os.path.join(RULES, fn), encoding='utf-8').read()
    doc = base_doc()
    render_md(doc, md)
    doc.add_paragraph()
    foot = doc.add_paragraph('CTRC — WonderKids · Quy tắc AI · cập nhật 2026-06-09 · nguồn: ctrc/rules/'+fn)
    foot.runs[0].font.size=Pt(8); foot.runs[0].font.color.rgb=RGBColor(0x99,0x99,0x99)
    out = os.path.join(OUT, title+'.docx'); doc.save(out)
    print('wrote', out)

# ---- bản gộp ----
doc = base_doc()
h = doc.add_heading('', level=0); r=h.add_run('CTRC — Bộ Quy tắc AI (Skill / Rules)');
sub = doc.add_paragraph('WonderKids · Tài liệu để khách hàng xem & chỉnh bổ sung nội dung AI viết báo cáo.')
sub.runs[0].italic=True
doc.add_paragraph('Gồm 4 bộ quy tắc AI dùng khi viết nội dung cho phụ huynh. Quý khách chỉnh trực tiếp trong file này; bên kỹ thuật sẽ cập nhật vào hệ thống. Cập nhật: 2026-06-09.')
for idx,(fn,title) in enumerate(FILES):
    doc.add_page_break()
    md = open(os.path.join(RULES, fn), encoding='utf-8').read()
    render_md(doc, md, as_section=True)
out = os.path.join(OUT, 'CTRC - Bo Quy tac AI (gop).docx'); doc.save(out)
print('wrote', out)
print('DONE -> '+OUT)
